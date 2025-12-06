"""
ATS View Generator - Uses PyMuPDF for accurate text extraction and layout analysis.

This is SEPARATE from the structured parser and always uses PyMuPDF to show
exactly how ATS systems will read the resume.
"""

import fitz  # PyMuPDF
from docx import Document
from typing import Dict, Any, List, Tuple
from collections import Counter
from app.services.ats_issues import compute_complexity_metric, compute_secondary_column_ratio


class ATSViewGenerator:
    """Generate ATS-compatible plain text view and extract layout diagnostics"""
    
    def generate_ats_view(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Generate ATS plain-text view and layout diagnostics.
        
        Returns:
            {
                "ats_text": str - Plain text as ATS would see it
                "diagnostics": {
                    "has_images": bool,
                    "has_tables": bool,
                    "has_headers_footers": bool,
                    "has_multi_column": bool,
                    "secondary_column_ratio": float,
                    "font_count": int,
                    "page_count": int,
                    "character_count": int,
                    "layout_complexity": str,  # simple/moderate/complex
                    "complexity_metric": dict,  # Detailed complexity breakdown
                    "warnings": List[str]
                }
            }
        """
        if file_type == ".pdf":
            return self._generate_from_pdf(file_path)
        elif file_type == ".docx":
            return self._generate_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _generate_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Generate ATS view and diagnostics from PDF"""
        doc = fitz.open(file_path)
        
        # Extract plain text (as ATS sees it)
        ats_text = ""
        for page in doc:
            ats_text += page.get_text("text")
        
        # Run diagnostics
        diagnostics = self._analyze_pdf_layout(doc)
        
        doc.close()
        
        return {
            "ats_text": ats_text,
            "diagnostics": diagnostics
        }
    
    def _generate_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Generate ATS view and diagnostics from DOCX"""
        doc = Document(file_path)
        
        # Extract plain text
        ats_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Basic diagnostics for DOCX
        diagnostics = self._analyze_docx_layout(doc)
        
        return {
            "ats_text": ats_text,
            "diagnostics": diagnostics
        }
    
    def _analyze_pdf_layout(self, doc: fitz.Document) -> Dict[str, Any]:
        """Analyze PDF layout for ATS compatibility issues"""
        has_images = False
        image_count = 0
        has_tables = False
        table_count = 0
        font_set = set()
        warnings = []
        all_blocks = []  # Collect blocks for column analysis
        
        for page_num, page in enumerate(doc):
            rect = page.rect
            page_width = rect.width
            
            # Check for images
            images = page.get_images()
            if images:
                has_images = True
                image_count += len(images)
                warnings.append(f"Page {page_num + 1}: Contains {len(images)} image(s) - ATS cannot read these")
            
            # Check for tables
            tables = page.find_tables()
            if tables.tables:
                has_tables = True
                table_count += len(tables.tables)
                warnings.append(f"Page {page_num + 1}: Contains tables - may confuse ATS systems")
            
            # Analyze fonts and extract blocks for column detection
            text_dict = page.get_text("dict")
            for block in text_dict["blocks"]:
                if block["type"] == 0:  # Text block
                    # Collect fonts
                    if "lines" in block:
                        text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                font_set.add(span["font"])
                                text += span.get("text", "") + " "
                        
                        # Store block for column analysis
                        if text.strip():
                            all_blocks.append({
                                "text": text.strip(),
                                "bbox": block["bbox"],
                                "page_width": page_width
                            })
        
        # Detect columns and compute secondary column ratio
        has_multi_column, secondary_column_ratio = self._detect_columns_and_ratio(all_blocks)
        
        # Compute detailed complexity metrics using new system
        font_count = len(font_set)
        
        # Check for headers/footers (heuristic: text in top/bottom 10% of page)
        has_headers_footers = self._detect_headers_footers(doc)
        if has_headers_footers:
            warnings.append("Headers/footers detected - may interfere with ATS parsing")
        
        # Compute complexity with all detected features
        complexity_metric = compute_complexity_metric(
            font_count=font_count,
            has_images=has_images,
            image_count=image_count,
            has_tables=has_tables,
            table_count=table_count,
            has_multi_column=has_multi_column,
            has_headers_footers=has_headers_footers,
            secondary_column_ratio=secondary_column_ratio
        )
        
        return {
            "has_images": has_images,
            "image_count": image_count,
            "has_tables": has_tables,
            "table_count": table_count,
            "has_multi_column": has_multi_column,
            "secondary_column_ratio": secondary_column_ratio,
            "has_headers_footers": has_headers_footers,
            "font_count": font_count,
            "page_count": len(doc),
            "character_count": len("".join([page.get_text() for page in doc])),
            "layout_complexity": complexity_metric.label.value,
            "complexity_metric": complexity_metric.to_dict(),
            "warnings": warnings
        }
    
    def _analyze_docx_layout(self, doc: Document) -> Dict[str, Any]:
        """Analyze DOCX layout for ATS compatibility"""
        warnings = []
        
        # Check for tables
        has_tables = len(doc.tables) > 0
        if has_tables:
            warnings.append(f"Document contains {len(doc.tables)} table(s) - may confuse ATS")
        
        # Check for images (inline shapes)
        has_images = False
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                has_images = True
                warnings.append("Document contains images - ATS cannot read these")
                break
        
        # Count paragraphs and text
        paragraph_count = len(doc.paragraphs)
        text = "\n".join([p.text for p in doc.paragraphs])
        
        # Determine complexity
        if not has_images and not has_tables and paragraph_count < 50:
            complexity = "simple"
        elif has_tables or paragraph_count > 100:
            complexity = "moderate"
        else:
            complexity = "simple"
        
        return {
            "has_images": has_images,
            "has_tables": has_tables,
            "has_headers_footers": False,  # Hard to detect in DOCX
            "has_multi_column": False,  # Hard to detect in DOCX
            "secondary_column_ratio": 0.0,
            "font_count": 0,  # Not easily accessible in python-docx
            "page_count": 1,  # Estimate
            "character_count": len(text),
            "layout_complexity": complexity,
            "warnings": warnings
        }
    
    def _detect_headers_footers(self, doc: fitz.Document) -> bool:
        """
        Detect if document has actual headers/footers.
        
        Smarter detection that avoids false positives:
        1. Uses smaller margin threshold (5% / ~36pt instead of 10%)
        2. For multi-page docs: checks if content REPEATS across pages
        3. For single-page docs: only flags if content is in extreme margins
           AND looks like typical header/footer content
        
        Returns:
            True if actual headers/footers detected, False otherwise
        """
        page_count = len(doc)
        
        if page_count == 0:
            return False
        
        # Use 5% margin (~36pt for letter-size) - where real headers/footers live
        # Not 10% which catches the main content (name, contact)
        MARGIN_RATIO = 0.05
        
        # For multi-page documents: check if content repeats
        if page_count > 1:
            return self._detect_repeating_headers_footers(doc, MARGIN_RATIO)
        
        # For single-page documents: use stricter criteria
        return self._detect_single_page_headers_footers(doc, MARGIN_RATIO)
    
    def _detect_repeating_headers_footers(self, doc: fitz.Document, margin_ratio: float) -> bool:
        """
        For multi-page documents, detect if the SAME content repeats
        in header/footer regions across pages.
        
        Real headers/footers repeat (e.g., "John Doe - Page 1", "John Doe - Page 2")
        """
        header_texts = []
        footer_texts = []
        
        for page in doc:
            rect = page.rect
            height = rect.height
            
            # Get text in header/footer margins
            top_rect = fitz.Rect(0, 0, rect.width, height * margin_ratio)
            bottom_rect = fitz.Rect(0, height * (1 - margin_ratio), rect.width, height)
            
            top_text = page.get_text("text", clip=top_rect).strip()
            bottom_text = page.get_text("text", clip=bottom_rect).strip()
            
            if top_text:
                # Normalize: remove page numbers for comparison
                normalized = self._normalize_header_footer_text(top_text)
                header_texts.append(normalized)
            
            if bottom_text:
                normalized = self._normalize_header_footer_text(bottom_text)
                footer_texts.append(normalized)
        
        # Check if headers repeat (same content on 2+ pages)
        if len(header_texts) >= 2:
            # Check if any header text appears multiple times
            if self._has_repeating_content(header_texts):
                return True
        
        # Check if footers repeat
        if len(footer_texts) >= 2:
            if self._has_repeating_content(footer_texts):
                return True
        
        return False
    
    def _detect_single_page_headers_footers(self, doc: fitz.Document, margin_ratio: float) -> bool:
        """
        For single-page documents, only flag if content is in extreme margins
        AND looks like typical header/footer content (not main resume content).
        
        We're stricter here because most single-page resumes start with
        name/contact at the top, which is NOT a header.
        """
        if len(doc) == 0:
            return False
        
        page = doc[0]
        rect = page.rect
        height = rect.height
        
        # Use even smaller margin for single-page (3% / ~24pt)
        # This is the extreme top/bottom where only true headers/footers live
        STRICT_MARGIN = 0.03
        
        bottom_rect = fitz.Rect(0, height * (1 - STRICT_MARGIN), rect.width, height)
        bottom_text = page.get_text("text", clip=bottom_rect).strip()
        
        # Only check footer for single page (top always has name/contact)
        if bottom_text:
            # Check if it looks like footer content
            if self._looks_like_footer_content(bottom_text):
                return True
        
        return False
    
    def _normalize_header_footer_text(self, text: str) -> str:
        """
        Normalize text for comparison by removing page numbers and whitespace.
        This helps detect repeating content like "John Doe Resume - Page X"
        """
        import re
        # Remove page numbers like "Page 1", "1/3", "- 1 -", etc.
        text = re.sub(r'page\s*\d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\d+\s*/\s*\d+', '', text)
        text = re.sub(r'-\s*\d+\s*-', '', text)
        text = re.sub(r'\s+', ' ', text).strip()
        return text.lower()
    
    def _has_repeating_content(self, texts: List[str]) -> bool:
        """Check if any content repeats in the list (appears 2+ times)"""
        if len(texts) < 2:
            return False
        
        # Count occurrences
        from collections import Counter
        counts = Counter(texts)
        
        # If any text appears on 2+ pages, it's likely a header/footer
        for text, count in counts.items():
            if text and count >= 2:
                return True
        
        return False
    
    def _looks_like_footer_content(self, text: str) -> bool:
        """
        Check if text looks like typical footer content.
        
        Common footer patterns:
        - Page numbers: "Page 1", "1/3", "1 of 3"
        - Confidential notices
        - Dates
        - "Continued on next page"
        """
        import re
        text_lower = text.lower()
        
        # Page number patterns
        if re.search(r'page\s*\d+', text_lower):
            return True
        if re.search(r'\d+\s*/\s*\d+', text_lower):  # "1/3"
            return True
        if re.search(r'\d+\s+of\s+\d+', text_lower):  # "1 of 3"
            return True
        if re.search(r'-\s*\d+\s*-', text_lower):  # "- 1 -"
            return True
        
        # Common footer keywords
        footer_keywords = [
            'confidential', 'continued', 'page', 'resume of',
            'curriculum vitae', 'cv of'
        ]
        for keyword in footer_keywords:
            if keyword in text_lower:
                return True
        
        return False
    
    def _detect_columns_and_ratio(self, blocks: List[Dict[str, Any]]) -> Tuple[bool, float]:
        """
        Detect multi-column layout and compute secondary column ratio.
        
        Args:
            blocks: List of text blocks with bbox and text
        
        Returns:
            Tuple of (has_multi_column: bool, secondary_column_ratio: float)
        """
        if not blocks:
            return False, 0.0
        
        # Simple column detection based on x-position clustering
        # Get the left edge (x0) of each block
        x_positions = [block["bbox"][0] for block in blocks]
        page_width = blocks[0].get("page_width", 612.0) if blocks else 612.0
        
        # Sort x positions to find gaps
        x_sorted = sorted(set(x_positions))
        
        # Look for significant gaps (> 20% of page width)
        gap_threshold = page_width * 0.2
        column_boundaries = [x_sorted[0]]

        for i in range(1, len(x_sorted)):
            if x_sorted[i] - x_sorted[i-1] > gap_threshold:
                column_boundaries.append(x_sorted[i])

        # Assign column numbers to blocks
        blocks_with_columns = []
        for block in blocks:
            x0 = block["bbox"][0]
            # Find which column this block belongs to
            column = 1
            for i, boundary in enumerate(column_boundaries):
                if x0 >= boundary:
                    column = i + 1

            blocks_with_columns.append({
                "text": block["text"],
                "column": column
            })

        # Compute secondary column ratio using the helper function
        secondary_ratio = compute_secondary_column_ratio(blocks_with_columns)
        
        # Only consider it multi-column if there's actually content in secondary columns
        # (Multiple boundaries might exist, but if all content is in column 1, it's not multi-column)
        has_multi_column = len(column_boundaries) > 1 and secondary_ratio > 0.0
        
        # Debug output
        if len(column_boundaries) > 1:
            print(f"\n[COLUMN DETECTION DEBUG]")
            print(f"  Column boundaries found: {len(column_boundaries)} (at positions: {[f'{b:.1f}' for b in column_boundaries]})")
            print(f"  Secondary column ratio: {secondary_ratio:.2%}")
            print(f"  Has multi-column (after ratio check): {has_multi_column}")
            # Show column distribution
            col_counts = Counter([b.get('column', 1) for b in blocks_with_columns])
            print(f"  Blocks per column: {dict(col_counts)}")
            print()

        return has_multi_column, secondary_ratio
    
    def extract_text_with_coordinates(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract text with bounding box coordinates (useful for advanced features).
        
        Returns list of text blocks with coordinates:
        [
            {
                "text": str,
                "x0": float, "y0": float, "x1": float, "y1": float,
                "page": int,
                "font_size": float,
                "font_name": str
            }
        ]
        """
        doc = fitz.open(file_path)
        blocks = []
        
        for page_num, page in enumerate(doc):
            text_dict = page.get_text("dict")
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            blocks.append({
                                "text": span["text"],
                                "x0": span["bbox"][0],
                                "y0": span["bbox"][1],
                                "x1": span["bbox"][2],
                                "y1": span["bbox"][3],
                                "page": page_num + 1,
                                "font_size": span["size"],
                                "font_name": span["font"]
                            })
        
        doc.close()
        return blocks
    